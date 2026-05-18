from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document

from utils.achievement_report import replace_text, set_font


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_APPLICATION_TEMPLATE_PATH = PROJECT_ROOT / "assets" / "活動申請書模板.docx"


def roc_date_from_iso(value: str) -> str:
    text = str(value).strip()
    if not text:
        return ""

    try:
        parsed = datetime.strptime(text, "%Y-%m-%d").date()
    except ValueError:
        return text

    return f"{parsed.year - 1911}/{parsed.month}/{parsed.day}"


def safe_int(value: object, default: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        set_font(paragraph)


def officer_names_by_role(officers: list[dict[str, str]], role: str) -> str:
    names = [
        officer.get("姓名", "").strip()
        for officer in officers
        if officer.get("職位", "").strip() == role and officer.get("姓名", "").strip()
    ]
    return " / ".join(names)


def first_officer_name(officers: list[dict[str, str]]) -> str:
    for officer in officers:
        name = officer.get("姓名", "").strip()
        if name:
            return name
    return ""


def update_work_assignments(
    doc,
    *,
    activity_leader: str,
    officers: list[dict[str, str]],
) -> None:
    if len(doc.tables) < 3:
        return

    table = doc.tables[2]
    president = officer_names_by_role(officers, "社長")
    vice_president = officer_names_by_role(officers, "副社長")
    secretary = officer_names_by_role(officers, "文書")
    photographer = officer_names_by_role(officers, "攝錄")
    treasurer = officer_names_by_role(officers, "總務")
    snack = officer_names_by_role(officers, "點心")

    leader = activity_leader.strip() or president or first_officer_name(officers)
    assignments = {
        2: leader,
        3: vice_president,
        4: president,
        5: vice_president,
        6: photographer,
        7: secretary,
        8: photographer,
        9: treasurer,
        10: snack,
    }

    for row_index, name in assignments.items():
        if name and row_index < len(table.rows) and len(table.rows[row_index].cells) > 1:
            set_cell_text(table.rows[row_index].cells[1], name)


def update_schedule(doc, *, activity_date: str, tea_topic: str) -> None:
    if len(doc.tables) < 4:
        return

    table = doc.tables[3]
    schedule_date = roc_date_from_iso(activity_date.split()[0] if activity_date else "")
    if schedule_date:
        for row_index in range(2, min(6, len(table.rows))):
            if table.rows[row_index].cells:
                set_cell_text(table.rows[row_index].cells[0], schedule_date)

    if tea_topic:
        replace_text(doc, {"{{茶 }}": tea_topic, "{{茶}}": tea_topic})


def update_budget(
    doc,
    *,
    estimated_people: int,
    snack_unit_price: int,
    snack_purpose: str,
) -> None:
    total = estimated_people * snack_unit_price

    if len(doc.tables) >= 5:
        equipment_table = doc.tables[4]
        if len(equipment_table.rows) > 8 and len(equipment_table.rows[8].cells) > 1:
            set_cell_text(equipment_table.rows[8].cells[1], str(estimated_people))

    if len(doc.tables) >= 6:
        budget_table = doc.tables[5]
        if len(budget_table.rows) > 2:
            row = budget_table.rows[2]
            values = {
                2: snack_unit_price,
                3: estimated_people,
                5: total,
                6: snack_purpose,
            }
            for cell_index, value in values.items():
                if cell_index < len(row.cells):
                    set_cell_text(row.cells[cell_index], str(value))

        if len(budget_table.rows) > 3:
            row = budget_table.rows[3]
            for cell_index in (5, 6, 7):
                if cell_index < len(row.cells):
                    set_cell_text(row.cells[cell_index], str(total))

    if len(doc.tables) >= 7:
        source_table = doc.tables[6]
        for row_index in (2, 7):
            if len(source_table.rows) > row_index and len(source_table.rows[row_index].cells) > 2:
                set_cell_text(source_table.rows[row_index].cells[2], str(total))


def build_application_form(
    *,
    template_file,
    fields: dict[str, object],
    officers: list[dict[str, str]],
) -> BytesIO:
    template_source = (
        template_file
        if template_file is not None
        else str(DEFAULT_APPLICATION_TEMPLATE_PATH)
    )
    doc = Document(template_source)

    activity_place = str(fields.get("activity_place", "")).strip()
    replacements = {
        "{{活動名稱}}": str(fields.get("activity_name", "")),
        "{{活動日期}}": str(fields.get("activity_date", "")),
        "{{活動負責人}}": str(fields.get("activity_leader", "")),
        "{{負責人電話}}": str(fields.get("leader_phone", "")),
        "{{活動宗旨}}": str(fields.get("activity_purpose", "")),
        "{{點心}}": str(fields.get("snack_purpose", "")),
    }
    if activity_place:
        replacements["靜心書院A202"] = activity_place

    replace_text(doc, replacements)
    update_work_assignments(
        doc,
        activity_leader=str(fields.get("activity_leader", "")),
        officers=officers,
    )
    update_schedule(
        doc,
        activity_date=str(fields.get("activity_date", "")),
        tea_topic=str(fields.get("tea_topic", "")),
    )
    update_budget(
        doc,
        estimated_people=safe_int(fields.get("estimated_people"), 30),
        snack_unit_price=safe_int(fields.get("snack_unit_price"), 100),
        snack_purpose=str(fields.get("snack_purpose", "")),
    )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
